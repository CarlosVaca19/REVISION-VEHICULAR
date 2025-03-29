from flask import Flask, request, render_template, send_file
from docx import Document
from io import BytesIO

app = Flask(__name__)

# Ruta del archivo modelo
MODEL_FILE = 'modelos/modelo.docx'

@app.route('/')
def index():
    # Cargar el archivo modelo
    doc = Document(MODEL_FILE)
    highlighted_text = []

    # Buscar texto resaltado
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color:  # Verifica si el texto está resaltado
                highlighted_text.append(run.text)

    return render_template('index.html', highlighted_text=highlighted_text)

@app.route('/modify', methods=['POST'])
def modify():
    # Cargar el archivo modelo
    doc = Document(MODEL_FILE)
    replacements = request.form.to_dict()

    # Obtener el número de oficio del formulario
    numero_oficio = replacements.get("Número de Oficio", "oficio_sin_numero")

    # Reemplazar texto resaltado y eliminar el resaltado
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color and run.text in replacements:
                run.text = replacements[run.text]  # Reemplazar el texto
                run.font.highlight_color = None  # Eliminar el resaltado

    # Guardar el archivo modificado en memoria
    modified_file = BytesIO()
    doc.save(modified_file)
    modified_file.seek(0)

    # Descargar el archivo con el número de oficio como nombre
    return send_file(modified_file, as_attachment=True, download_name=f'{numero_oficio}.docx')

if __name__ == '__main__':
    app.run(debug=True)